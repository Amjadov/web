"""
Generate a professional, ATS-friendly .docx resume for Amjad M. Masoud, PMP.
Uses python-docx. Run: python generate_resume.py
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colour palette ──────────────────────────────────────────────
NAVY    = RGBColor(0x1B, 0x2A, 0x4A)   # dark navy for headings
ACCENT  = RGBColor(0x2E, 0x74, 0xB5)   # blue accent for name / lines
DARK    = RGBColor(0x33, 0x33, 0x33)    # body text
MEDIUM  = RGBColor(0x55, 0x55, 0x55)    # secondary text
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Page setup ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

style = doc.styles['Normal']
style.font.name  = 'Calibri'
style.font.size  = Pt(10.5)
style.font.color.rgb = DARK
style.paragraph_format.space_after  = Pt(2)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

# ── Helper functions ────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    """Set cell borders. Pass keys like top, bottom, start, end with dict values."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_horizontal_line():
    """Add a thin navy horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="1B2A4A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_section_heading(text):
    """ATS-friendly section heading: UPPERCASE, navy, with underline bar."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    # bottom border on paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="2E74B5"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_role_header(title, dates, company):
    """Job title (bold), dates right-aligned, company on same line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(1)

    # Use a tab stop near the right margin for dates
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = NAVY
    run_title.font.name = 'Calibri'

    p.add_run('\t')

    run_dates = p.add_run(dates)
    run_dates.font.size = Pt(10)
    run_dates.font.color.rgb = MEDIUM
    run_dates.font.name = 'Calibri'

    # Company line
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(3)
    run_co = p2.add_run(company)
    run_co.italic = True
    run_co.font.size = Pt(10)
    run_co.font.color.rgb = ACCENT
    run_co.font.name = 'Calibri'
    return p2


def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
    return p


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.text = ''
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(1.0)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    return p


def add_tech_scope(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = MEDIUM
    run.font.name = 'Calibri'
    return p


# ═══════════════════════════════════════════════════════════════
# HEADER – Name & Contact
# ═══════════════════════════════════════════════════════════════

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_after = Pt(2)
run = name_p.add_run('AMJAD M. MASOUD, PMP')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(1)
run = contact_p.add_run('Amman, Jordan  ∙  amjadmasoud@hotmail.com  ∙  (+962) 799 312 787')
run.font.size = Pt(10)
run.font.color.rgb = MEDIUM
run.font.name = 'Calibri'

links_p = doc.add_paragraph()
links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
links_p.paragraph_format.space_after = Pt(4)
run = links_p.add_run('linkedin.com/in/amjad-m-masoud-pmp-9308aa11  ∙  amjadov.github.io/web')
run.font.size = Pt(9.5)
run.font.color.rgb = ACCENT
run.font.name = 'Calibri'

# Headline
hl = doc.add_paragraph()
hl.alignment = WD_ALIGN_PARAGRAPH.CENTER
hl.paragraph_format.space_before = Pt(2)
hl.paragraph_format.space_after  = Pt(4)
run = hl.add_run('Senior IT Leader  |  Project & Program Manager (PMP)  |  Enterprise Architecture (TOGAF)  |  Digital Transformation Strategist')
run.bold = True
run.font.size = Pt(10.5)
run.font.color.rgb = ACCENT
run.font.name = 'Calibri'

add_horizontal_line()

# ═══════════════════════════════════════════════════════════════
# PROFESSIONAL SUMMARY
# ═══════════════════════════════════════════════════════════════

add_section_heading('Professional Summary')

add_body(
    'Results-driven IT leader and PMP-certified Project Manager with 20+ years of experience '
    'steering large-scale digital transformation programs, enterprise modernization, and '
    'cross-functional team leadership. Proven ability to bridge the gap between C-suite strategy '
    'and technical execution — turning complex, legacy-laden environments into scalable, '
    'high-performance platforms.'
)
add_body(
    'Currently driving government process automation at EY and pursuing TOGAF Foundation '
    'certification to deepen enterprise architecture expertise. Track record of managing '
    'end-to-end project lifecycles, building high-performing teams, and delivering measurable '
    'business outcomes: 8× throughput gains, sub-200ms response times, and zero-downtime migrations.'
)
add_body(
    'Seeking senior leadership roles — Project/Program Manager, Enterprise Architect, or Head of IT '
    '— where I can apply my blend of strategic vision, architectural thinking, and hands-on '
    'delivery experience to drive organizational transformation at scale.'
)

# ═══════════════════════════════════════════════════════════════
# CORE COMPETENCIES (ATS-friendly: plain text grid, not images)
# ═══════════════════════════════════════════════════════════════

add_section_heading('Core Competencies')

competencies = [
    ['Project & Program Management',  'Enterprise Architecture (TOGAF)',    'Digital Transformation Strategy'],
    ['Agile / Scrum Delivery',        'Stakeholder & Executive Engagement', 'Risk & Issue Management'],
    ['Team Building & Leadership',    'Budgeting & Resource Planning',      'Vendor & Contract Management'],
    ['Legacy System Modernization',   'Cloud & Hybrid Architecture',        'Business Process Re-engineering'],
    ['CI/CD & DevOps Practices',      'Requirements Analysis (BRD/FRD)',    'Change Management'],
]

table = doc.add_table(rows=len(competencies), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

for i, row_data in enumerate(competencies):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'
        # Light background for alternating rows
        if i % 2 == 0:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EDF2F9" w:val="clear"/>')
            cell._tc.get_or_add_tcPr().append(shading)
        # Remove borders for clean ATS-friendly look
        no_border = {"val": "none", "sz": "0", "color": "FFFFFF"}
        set_cell_border(cell, top=no_border, bottom=no_border, start=no_border, end=no_border)

# ═══════════════════════════════════════════════════════════════
# TECHNICAL PROFICIENCIES
# ═══════════════════════════════════════════════════════════════

add_section_heading('Technical Proficiencies')

tech_items = [
    ('Platforms & Infrastructure:', 'Microsoft Azure, Cloud & Hybrid Architectures, SQL Server (Always On), Redis, Memcache, MQ'),
    ('Development:',               'C# .NET, ASP.NET, RESTful APIs, Python (Django, Data Science), PHP (Laravel, CodeIgniter)'),
    ('Project Delivery:',          'Agile-Scrum, Sprint Planning, CI/CD Pipelines, SDLC, UAT'),
    ('Tools & Frameworks:',        'Git, Jira, Microsoft Project, ERP Systems, Applied AI'),
]

for label, value in tech_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    run_label = p.add_run(label + '  ')
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = NAVY
    run_label.font.name = 'Calibri'
    run_val = p.add_run(value)
    run_val.font.size = Pt(10)
    run_val.font.color.rgb = DARK
    run_val.font.name = 'Calibri'

# ═══════════════════════════════════════════════════════════════
# PROFESSIONAL EXPERIENCE
# ═══════════════════════════════════════════════════════════════

add_section_heading('Professional Experience')

# ── EY ──────────────────────────────────────────────────────────
add_role_header(
    'Project Manager | Government Process Automation',
    'Feb 2025 – Present',
    'EY, Amman, Jordan'
)
add_tech_scope('Client: Saudi Expro – Services Efficiency Review Department')
add_body(
    'Leading the end-to-end digital transformation of a government agency\'s manual workflows, '
    'replacing fragmented legacy systems with a centralized, scalable platform. Managing a '
    'cross-functional team of 5 and reporting directly to client executives.'
)
for bullet in [
    'Drove full project lifecycle from initiation through UAT, applying Agile-Scrum methodology with 2-week sprint cadences to deliver on time and within budget.',
    'Authored all key project artefacts — Business Requirements Document (BRD), user stories, interface mockups, and user manuals — ensuring traceability from requirements to delivered functionality.',
    'Digitized core business processes, eliminating spreadsheet/email-based workflows and establishing an automated system that improved data integrity and operational efficiency.',
    'Facilitated stakeholder workshops and weekly executive sprint reviews, securing formal client sign-off on all major deliverables.',
    'Maintained project health through proactive risk register management, sprint velocity tracking, and continuous alignment with budgetary and schedule baselines.',
]:
    add_bullet(bullet)

# ── MStart Lead ─────────────────────────────────────────────────
add_role_header(
    'Development Team Lead | Technical Architect',
    '2021 – 2024',
    'MStart, Amman, Jordan'
)
add_body(
    'Recruited to lead the technical overhaul of a high-traffic e-commerce platform, owning both '
    'team leadership and solution architecture. Built and managed dual-stream backend and mobile '
    'teams while driving a legacy-to-modern platform migration.'
)
for bullet in [
    'Scaled platform capacity from 15,000 to 120,000 daily orders (8× increase) while reducing response times from 3s to 200ms — directly enabling business revenue growth.',
    'Built and led a cross-functional team (Backend + Mobile), owning recruitment, onboarding, daily stand-ups, and blocker resolution to maintain sprint velocity.',
    'Architected SQL Server "Always On" high-availability solution, achieving zero-downtime during full platform migration.',
    'Established CI/CD pipelines that streamlined release cycles and reduced deployment risk.',
    'Directed the migration to a modern Android application, significantly reducing bug reports and improving end-user experience.',
    'Balanced hands-on technical delivery (C#, SQL) with people management and stakeholder communication.',
]:
    add_bullet(bullet)

# ── MStart Senior ───────────────────────────────────────────────
add_role_header(
    'Senior Systems Developer',
    '2018 – 2021',
    'MStart, Amman, Jordan'
)
add_tech_scope('C# .NET, RESTful APIs, SQL Server, Redis, Memcache, MQ Services')
add_body(
    'Led task allocation and issue resolution across the development team while driving '
    'architectural improvements to core platform services.'
)
for bullet in [
    'Re-architected SOAP services into lightweight RESTful APIs, improving maintainability and enabling the platform to process 200,000+ transactions per day — a 400% increase in sales throughput.',
    'Diagnosed and resolved critical performance bottlenecks (deadlocks, timeouts) achieving millisecond-level transaction improvements and enabling horizontal scale-out of applications and databases.',
    'Designed and implemented optimized stored procedures and user-defined data types to support high-volume transactional workloads.',
]:
    add_bullet(bullet)

# ── Kafou ───────────────────────────────────────────────────────
add_role_header(
    'Software Development Consultant',
    '2017 – Present',
    'Kafou Services, Jordan'
)
add_tech_scope('C# ASP.NET, Python Django, PHP Laravel/CodeIgniter, SQL Server, MySQL, Scikit-Learn, Twilio SDK')
add_body(
    'Provide ongoing technical consulting across multiple concurrent projects, overseeing '
    'architecture decisions, development standards, and delivery timelines.'
)
for bullet in [
    'Led machine learning research initiatives using Python and Scikit-Learn for predictive analytics.',
    'Architected two-way data synchronization solution using Python Django, Python Signals, and MySQL.',
    'Delivered "ComplianceArch" web application (C# ASP.NET, SQL Server) for regulatory compliance management.',
    'Managed end-to-end delivery of Kafou backend services and admin panel (PHP CodeIgniter, MySQL).',
    'Designed Wayt platform API with online payment integration (PayFort), real-time messaging (PubNub), and geolocation (Google Maps).',
    'Implemented peer-to-peer and group video calling solution using PHP Laravel, MySQL, and Twilio SDK.',
]:
    add_bullet(bullet)

# ── Abbar ERP ───────────────────────────────────────────────────
add_role_header(
    'ERP Implementation Specialist',
    '2013 – 2017',
    'Mahmoud Saleh Abbar, Jeddah, Saudi Arabia'
)
add_tech_scope('SQL Server, .NET Applications, ERP Systems')
add_body(
    'Led a company-wide ERP implementation across 5 cities and 450+ employees, while developing '
    'supporting .NET applications for management reporting.'
)
for bullet in [
    'Drove enterprise-wide ERP rollout, coordinating across multiple locations and stakeholder groups to ensure adoption and data migration integrity.',
    'Automated back-office functions spanning technology, services, and human resources — streamlining business processes across the organization.',
    'Designed and implemented advanced T-SQL stored procedures to produce data pipelines for statistics, finance, and auditing.',
]:
    add_bullet(bullet)

# ── Grant Thornton ME ───────────────────────────────────────────
add_role_header(
    'Senior Full Stack Developer',
    '2010 – 2013',
    'Grant Thornton ME, Amman, Jordan'
)
add_tech_scope('C# .NET, ASP.NET, SQL Server, JavaScript, Telerik')
add_body(
    'Owned development and delivery of large-scale .NET applications deployed across the Middle East region.'
)
for bullet in [
    'Led the implementation of enterprise .NET applications across multiple countries in the Middle East.',
    'Drove development of new application modules, managing requirements gathering through deployment and user training.',
]:
    add_bullet(bullet)

# ── Grant Thornton SA ───────────────────────────────────────────
add_role_header(
    'Senior Full Stack Developer',
    '2005 – 2010',
    'Grant Thornton, Jeddah, Saudi Arabia'
)
add_tech_scope('C# .NET, ASP.NET, SQL Server, JavaScript, Telerik')
add_body(
    'Managed a development team of 2, coordinating with Grant Thornton SA and CGC UK leadership. '
    'Owned client relationship management, technical delivery, and cross-border collaboration with '
    'frequent travel to the UK.'
)
for bullet in [
    'Led development of an in-house software product that was successfully deployed at 25+ banks and financial institutions across the region — demonstrating product management and scaled delivery capability.',
]:
    add_bullet(bullet)

# Additional experience note
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run = p.add_run('Additional experience as Full Stack Developer at Visual Soft, Jubail, Saudi Arabia')
run.italic = True
run.font.size = Pt(9.5)
run.font.color.rgb = MEDIUM
run.font.name = 'Calibri'

# ═══════════════════════════════════════════════════════════════
# EDUCATION & CREDENTIALS
# ═══════════════════════════════════════════════════════════════

add_section_heading('Education')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run = p.add_run('Bachelor\'s Degree in Business Administration')
run.bold = True
run.font.size = Pt(10.5)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
run = p2.add_run('Applied Science University, Amman, Jordan')
run.font.size = Pt(10)
run.font.color.rgb = DARK
run.font.name = 'Calibri'

# ═══════════════════════════════════════════════════════════════
# CERTIFICATIONS
# ═══════════════════════════════════════════════════════════════

add_section_heading('Professional Certifications')

certs = [
    ('Project Management Professional (PMP®)',                'Project Management Institute (PMI)',      'March 2026 – March 2029'),
    ('TOGAF Foundation (Level 1)',                             '',                                        'In Progress'),
    ('Microsoft Certified: Azure Fundamentals (AZ-900)',      'Microsoft',                               'November 2025'),
    ('EY Applied AI – Bronze Learning',                       'EY',                                      'March 2025'),
    ('Building Business Acumen',                              'EY',                                      'March 2026'),
    ('Developing Client Relationships',                       'EY',                                      'March 2026'),
    ('Applied Data Science with Python – Level 2',            'IBM',                                     '2017'),
    ('Python for Data Science',                               'IBM',                                     '2017'),
    ('Advanced Certified Operational Risk Executive (CORE2)', 'Corporate Governance Consultancy',        '2010'),
    ('Certified Operational Risk Executive (CORE)',           'Aldar Audit Bureau / Grant Thornton',     'April 2007'),
]

for cert_name, issuer, date in certs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(1)

    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    run_name = p.add_run(cert_name)
    run_name.bold = True
    run_name.font.size = Pt(10)
    run_name.font.color.rgb = DARK
    run_name.font.name = 'Calibri'

    if issuer:
        run_issuer = p.add_run(f'  |  {issuer}')
        run_issuer.font.size = Pt(10)
        run_issuer.font.color.rgb = MEDIUM
        run_issuer.font.name = 'Calibri'

    p.add_run('\t')

    run_date = p.add_run(date)
    run_date.font.size = Pt(9.5)
    run_date.font.color.rgb = MEDIUM
    run_date.font.name = 'Calibri'

# ═══════════════════════════════════════════════════════════════
# LANGUAGES
# ═══════════════════════════════════════════════════════════════

add_section_heading('Languages')

langs = [
    ('Arabic', 'Native'),
    ('English', 'Fluent (Professional Working Proficiency)'),
    ('French', 'Intermediate'),
]

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
parts = []
for lang, level in langs:
    run_lang = p.add_run(lang + ':  ')
    run_lang.bold = True
    run_lang.font.size = Pt(10)
    run_lang.font.color.rgb = NAVY
    run_lang.font.name = 'Calibri'
    run_level = p.add_run(level)
    run_level.font.size = Pt(10)
    run_level.font.color.rgb = DARK
    run_level.font.name = 'Calibri'
    if lang != 'French':
        sep = p.add_run('     ∙     ')
        sep.font.size = Pt(10)
        sep.font.color.rgb = MEDIUM

# ═══════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), 'Amjad_Masoud_PMP_Resume.docx')
doc.save(output_path)
print(f'✅ Resume saved to: {output_path}')
